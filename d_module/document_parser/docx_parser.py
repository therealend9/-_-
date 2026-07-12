from __future__ import annotations
"""DOCX 解析工具。

当前策略：
1. DOCX 可编辑文本直接提取，避免 OCR 造成文字误差。
2. 遇到内嵌图片时，把图片导出到 storage/docx_images/。
3. 输出按文档出现顺序排列的 mixed elements，供 ocr_engine 生成 OCRPageResult。

依赖：pip install python-docx
"""

from pathlib import Path
from typing import Dict, Iterable, List, Optional, Union

from d_module import config


def _require_docx():
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("DOCX 解析需要安装 python-docx：pip install python-docx") from exc
    return docx


def _safe_image_ext(partname: object, content_type: Optional[str] = None) -> str:
    suffix = Path(str(partname)).suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"}:
        return ".jpg" if suffix == ".jpeg" else suffix

    if content_type:
        mapping = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/bmp": ".bmp",
            "image/gif": ".gif",
            "image/tiff": ".tif",
            "image/webp": ".webp",
        }
        return mapping.get(content_type.lower(), ".png")

    return ".png"


def _find_image_rel_ids(run) -> List[str]:
    """从一个 run 中找出内嵌图片 relationship id。"""
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return []

    rel_ids: List[str] = []
    blips = run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
    for blip in blips:
        rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rel_id:
            rel_ids.append(rel_id)
    return rel_ids


def _save_related_image(document, rel_id: str, file_id: str, image_index: int) -> Dict:
    image_part = document.part.related_parts.get(rel_id)
    if image_part is None:
        raise ValueError(f"DOCX 图片关系不存在: {rel_id}")

    ext = _safe_image_ext(
        getattr(image_part, "partname", ""),
        getattr(image_part, "content_type", None),
    )
    image_id = f"img_{image_index:03d}"
    output_path = config.EXTRACTED_IMAGE_DIR / f"{file_id}_{image_id}{ext}"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(image_part.blob)

    return {
        "type": "image",
        "image_id": image_id,
        "image_path": str(output_path),
        "rel_id": rel_id,
    }


def _paragraph_to_elements(paragraph, document, file_id: str, image_state: Dict[str, int]) -> List[Dict]:
    """把段落拆成文本元素和图片元素，尽量保持 run 内出现顺序。"""
    elements: List[Dict] = []
    text_buffer: List[str] = []

    def flush_text() -> None:
        text = "".join(text_buffer).strip()
        if text:
            elements.append({"type": "text", "text": text})
        text_buffer.clear()

    for run in paragraph.runs:
        if run.text:
            text_buffer.append(run.text)

        rel_ids = _find_image_rel_ids(run)
        for rel_id in rel_ids:
            flush_text()
            image_state["index"] += 1
            try:
                elements.append(_save_related_image(document, rel_id, file_id, image_state["index"]))
            except Exception as exc:
                # 图片提取失败时仍保留一个标记，避免文档顺序信息完全丢失。
                image_id = f"img_{image_state['index']:03d}"
                elements.append({
                    "type": "image",
                    "image_id": image_id,
                    "image_path": "",
                    "extract_error": str(exc),
                })

    flush_text()
    return elements


def _iter_block_items(document) -> Iterable[object]:
    """按正文顺序遍历段落和表格。"""
    try:
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        return []

    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def extract_docx_mixed_elements(docx_path: Union[str, Path], file_id: str) -> List[Dict]:
    """提取 DOCX 文本和内嵌图片，保持它们在文档中的出现顺序。

    返回示例：
    [
      {"type": "text", "text": "1. 什么是操作系统？"},
      {"type": "image", "image_id": "img_001", "image_path": ".../file_img_001.png"},
      {"type": "text", "text": "我的答案是..."}
    ]
    """
    config.ensure_storage_dirs()
    docx = _require_docx()
    document = docx.Document(str(docx_path))

    elements: List[Dict] = []
    image_state = {"index": 0}

    for block in _iter_block_items(document):
        # 段落：保留 run 中文本和图片的相对顺序。
        if hasattr(block, "runs"):
            elements.extend(_paragraph_to_elements(block, document, file_id, image_state))
            continue

        # 表格：第一版把每行表格文字拼成一个 text block；表格内图片暂按单元格段落顺序追加。
        if hasattr(block, "rows"):
            for row in block.rows:
                row_text_parts = []
                row_image_elements: List[Dict] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text_parts.append(cell_text)
                    for paragraph in cell.paragraphs:
                        row_image_elements.extend(
                            element for element in _paragraph_to_elements(paragraph, document, file_id, image_state)
                            if element.get("type") == "image"
                        )
                if row_text_parts:
                    elements.append({"type": "text", "text": " | ".join(row_text_parts)})
                elements.extend(row_image_elements)

    return elements


def extract_docx_text(docx_path: Union[str, Path]) -> str:
    """兼容旧接口：只返回 DOCX 可编辑文本，不处理图片。"""
    docx = _require_docx()
    document = docx.Document(str(docx_path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(lines)


def get_docx_basic_info(docx_path: Union[str, Path]) -> Dict:
    text = extract_docx_text(docx_path)
    return {
        "page_count": 1,
        "paragraph_count": len([line for line in text.splitlines() if line.strip()]),
        "char_count": len(text),
    }


def extract_docx_text_pages(docx_path: Union[str, Path]) -> List[Dict]:
    """兼容旧接口：第一版把 DOCX 可编辑文本作为单页文本处理。"""
    return [{"page_no": 1, "text": extract_docx_text(docx_path)}]
